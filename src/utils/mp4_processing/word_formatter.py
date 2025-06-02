# utils/word_formatter.py
"""
Word Document Formatter
Converts markdown to professionally formatted Word documents
"""

import io
import os
import re
import logging
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.section import WD_ORIENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from htmldocx import HtmlToDocx
import mistune
from bs4 import BeautifulSoup

from ..logger_setup import setup_logger

logger = setup_logger(name = "WordDocFormatter", level=logging.INFO)

class WordDocFormatter:
    """
    Utility class for converting markdown to formatted Word documents.
    This class provides methods to create professionally formatted
    Word documents from markdown strings, with support for templates
    and placeholder replacement.

    Key Features:
    - Convert markdown to Word documents with custom styling
    - Replace placeholders in Word templates with values from dictionaries
    - Support for both simple text and markdown content in placeholders
    - Preserve or override template styles
    - Custom table formatting with professional styling
    - Header/footer placeholder replacement

    Example usage for markdown conversion:
        formatter = WordDocFormatter()
        result = formatter.markdown_to_docx(
            markdown_text=md_content,
            output_filepath="output.docx",
            reference_docx_path="template.docx",
            preserve_reference_styles=True,      # Keep template styles
            custom_table_style_override=True     # But apply custom table formatting
        )

    Example usage for placeholder replacement:
        formatter = WordDocFormatter()
        placeholders = {
            'company_name': 'Acme Corporation',
            'date': '2024-01-15',
            'summary': '# Executive Summary\\n\\nThis is our quarterly report...',
            'revenue': '$1,250,000'
        }
        result = formatter.replace_placeholders(
            template_path="template.docx",
            placeholders=placeholders,
            output_filepath="filled_template.docx",
            convert_markdown=True,              # Convert markdown in placeholder values
            preserve_template_styles=True       # Keep original template formatting
        )
    """

    def __init__(self,
                 font_name="Calibri",
                 font_size=10,
                 heading_font="Calibri",
                 heading_sizes={'h1': 16, 'h2': 14, 'h3': 12, 'h4': 11},
                 heading_color="#025B95",
                 table_header_bg_color="#025B95",
                 custom_table_style=True,
                 page_orientation="portrait",
                 margins={'top': 1, 'bottom': 1, 'left': 1, 'right': 1}):
        """
        Initialize the WordDocFormatter with default formatting options.
        
        Args:
            font_name (str): Default font name for body text
            font_size (int): Default font size in points for body text
            heading_font (str): Font name for headings
            heading_sizes (dict): Dictionary mapping heading levels to font sizes
            heading_color (str): Hex color code for headings
            table_header_bg_color (str): Hex color code for table header background
            custom_table_style (bool): Whether to apply custom table styling
            page_orientation (str): Page orientation ('portrait' or 'landscape')
            margins (dict): Page margins in inches for 'top', 'bottom', 'left', 'right'
        """
        self.font_name = font_name
        self.font_size = font_size
        self.heading_font = heading_font
        self.heading_sizes = heading_sizes
        self.heading_color = heading_color
        self.table_header_bg_color = table_header_bg_color
        self.custom_table_style = custom_table_style
        self.page_orientation = page_orientation
        self.margins = margins
        
        # Pre-compile regex patterns for better performance
        self._markdown_patterns = [
            re.compile(r'^#{1,6}\s+', re.MULTILINE),  # Headers
            re.compile(r'^\*\s+', re.MULTILINE),       # Unordered lists
            re.compile(r'^\d+\.\s+', re.MULTILINE),    # Ordered lists
            re.compile(r'\*\*[^*]+\*\*'),              # Bold
            re.compile(r'\*[^*]+\*'),                  # Italic
            re.compile(r'```[\s\S]*```'),              # Code blocks
            re.compile(r'\|.*\|.*\|'),                 # Tables
            re.compile(r'^\s*[-*+]\s+', re.MULTILINE), # Alternative list markers
            re.compile(r'\[.*\]\(.*\)'),               # Links
        ]
        
        # Configure logging
        self.logger = logging.getLogger(__name__)
    
    def format_table_with_custom_style(self, doc):
        """
        Apply custom table formatting:
        - Header: Bold, white text, specified font, custom background color
        - Body: Specified font, black text
        - Borders: Only top and bottom (and horizontal inside borders)
        - Font size: Consistent throughout table
        
        Args:
            doc (Document): The docx Document object to be modified
        """
        # Define XML for borders and cell margins
        borders_xml = f'''
        <w:tblBorders {nsdecls('w')}>
            <w:top w:val="single" w:sz="4" w:color="auto"/>
            <w:left w:val="nil"/>
            <w:bottom w:val="single" w:sz="4" w:color="auto"/>
            <w:right w:val="nil"/>
            <w:insideH w:val="single" w:sz="4" w:color="auto"/>
            <w:insideV w:val="nil"/>
        </w:tblBorders>
        '''
        margin_xml = f'''
        <w:tcMar {nsdecls('w')}>
            <w:top w:w="70" w:type="dxa"/>
            <w:bottom w:w="70" w:type="dxa"/>
            <w:left w:w="100" w:type="dxa"/>
            <w:right w:w="100" w:type="dxa"/>
        </w:tcMar>
        '''

        # Remove # from color hex if present
        header_bg_color = self.table_header_bg_color.lstrip('#')

        for table in doc.tables:
            # Apply border settings
            table._tblPr.append(parse_xml(borders_xml))
            for row_idx, row in enumerate(table.rows):
                for cell in row.cells:
                    # Set cell margins
                    cell._element.tcPr.append(parse_xml(margin_xml))
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.name = self.font_name
                            run.font.size = Pt(self.font_size)
                            if row_idx == 0:  # Header row styling
                                run.font.bold = True
                                run.font.color.rgb = RGBColor(255, 255, 255)  # White
                            else:
                                run.font.color.rgb = RGBColor(0, 0, 0)  # Black
                    # Apply header background shading
                    if row_idx == 0:
                        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{header_bg_color}"/>')
                        cell._tc.get_or_add_tcPr().append(shading)
    
    def _is_likely_markdown(self, text):
        """
        More sophisticated markdown detection.
        
        Args:
            text (str): Text to check for markdown patterns
            
        Returns:
            bool: True if text likely contains markdown
        """
        if not text or not isinstance(text, str):
            return False
            
        # Check for markdown patterns
        for pattern in self._markdown_patterns:
            if pattern.search(text):
                return True
        return False
    
    def _find_and_replace_in_paragraph(self, paragraph, find_text, replace_text):
        """
        Find and replace text in a paragraph while preserving formatting.
        
        Args:
            paragraph: The paragraph object
            find_text: Text to find
            replace_text: Text to replace with
        """
        if find_text in paragraph.text:
            # Collect all runs and their text
            runs = list(paragraph.runs)
            full_text = paragraph.text
            
            # Find all occurrences
            occurrences = []
            start = 0
            while True:
                pos = full_text.find(find_text, start)
                if pos == -1:
                    break
                occurrences.append(pos)
                start = pos + 1
            
            # Process each occurrence
            for occurrence in reversed(occurrences):  # Process in reverse to maintain positions
                # Clear paragraph
                for run in runs:
                    run.text = ""
                
                # Rebuild paragraph with replacement
                current_pos = 0
                for run in runs:
                    run_start = full_text.find(run.text, current_pos) if run.text else current_pos
                    run_end = run_start + len(run.text)
                    
                    # Check if this run contains part of the text to replace
                    if occurrence < run_end and occurrence + len(find_text) > run_start:
                        # This run contains part of the text to replace
                        before_text = full_text[run_start:occurrence] if occurrence > run_start else ""
                        after_text = full_text[occurrence + len(find_text):run_end] if occurrence + len(find_text) < run_end else ""
                        run.text = before_text + replace_text + after_text
                        full_text = full_text[:occurrence] + replace_text + full_text[occurrence + len(find_text):]
                    else:
                        run.text = full_text[run_start:run_end]
                    
                    current_pos = run_end
    
    def _find_and_replace_in_tables(self, doc, find_text, replace_text):
        """Find and replace text in all tables."""
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self._find_and_replace_in_paragraph(paragraph, find_text, replace_text)
    
    def _replace_placeholder_with_markdown(self, doc, placeholder, markdown_text):
        """
        Replace a placeholder with formatted markdown content.
        
        Args:
            doc: Document object
            placeholder: Placeholder text to find
            markdown_text: Markdown content to insert
            
        Returns:
            bool: True if replacement was made
        """
        # Convert markdown to HTML
        html = mistune.create_markdown(plugins=['table'], escape=False)(markdown_text)
        
        # Create a temporary document
        temp_doc = Document()
        
        # Add HTML to temporary document
        try:
            HtmlToDocx().add_html_to_document(html, temp_doc)
        except Exception as e:
            self.logger.warning(f"HtmlToDocx failed, using custom parser: {e}")
            self._custom_html_parser(html, temp_doc)
        
        # Apply formatting to temp doc elements
        for p in temp_doc.paragraphs:
            for run in p.runs:
                run.font.name = self.font_name
                run.font.size = Pt(self.font_size)
        
        # Find placeholder in document
        replaced = False
        
        # Check paragraphs
        for i, paragraph in enumerate(doc.paragraphs):
            if placeholder in paragraph.text:
                # Insert new content after this paragraph
                for element in reversed(temp_doc.element.body):
                    paragraph._element.addnext(element)
                
                # Remove the placeholder paragraph
                paragraph._element.getparent().remove(paragraph._element)
                replaced = True
                break
        
        # Check tables if not found in paragraphs
        if not replaced:
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for i, paragraph in enumerate(cell.paragraphs):
                            if placeholder in paragraph.text:
                                # Clear cell and add new content
                                cell._element.clear_content()
                                for element in temp_doc.element.body:
                                    cell._element.append(element)
                                replaced = True
                                break
                        if replaced:
                            break
                    if replaced:
                        break
                if replaced:
                    break
        
        return replaced
    
    def _ensure_required_styles(self, doc):
        """
        Ensure that required styles exist in the document.
        Creates them if they don't exist.
        """
        styles = doc.styles
        
        # Define required styles with proper types
        required_styles = {
            'List Number': ('List Paragraph', WD_STYLE_TYPE.PARAGRAPH),
            'List Bullet': ('List Paragraph', WD_STYLE_TYPE.PARAGRAPH),
            'List Paragraph': ('Normal', WD_STYLE_TYPE.PARAGRAPH),
        }
        
        existing_style_names = [s.name for s in styles]
        
        for style_name, (base_style_name, style_type) in required_styles.items():
            if style_name not in existing_style_names:
                try:
                    new_style = styles.add_style(style_name, style_type)
                    
                    # Set base style if it exists
                    if base_style_name in existing_style_names:
                        new_style.base_style = styles[base_style_name]
                    
                    # Configure list-specific properties
                    if 'List' in style_name:
                        new_style.paragraph_format.left_indent = Inches(0.5)
                        new_style.paragraph_format.first_line_indent = Inches(-0.25)
                        
                except Exception as e:
                    self.logger.warning(f"Could not create style '{style_name}': {e}")
    
    def _custom_html_parser(self, html_content, doc):
        """
        Custom HTML parser that handles missing styles gracefully.
        Falls back to basic paragraph styles when specific styles are missing.
        """
        soup = BeautifulSoup(html_content, 'html.parser')
        
        # Track list counters for proper numbering
        ol_counters = {}
        list_stack = []
        
        def process_element(element, parent_list_info=None):
            """Recursively process HTML elements."""
            
            if element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                # Handle headings
                level = int(element.name[1])
                p = doc.add_paragraph(element.get_text().strip())
                try:
                    p.style = f'Heading {level}'
                except KeyError:
                    # If heading style doesn't exist, format manually
                    p.style = 'Normal'
                    for run in p.runs:
                        run.font.size = Pt(self.heading_sizes.get(f'h{level}', 14))
                        run.font.bold = True
                        run.font.name = self.heading_font
                        if self.heading_color:
                            hex_color = self.heading_color.lstrip('#')
                            r, g, b = int(hex_color[:2], 16), int(hex_color[2:4], 16), int(hex_color[4:6], 16)
                            run.font.color.rgb = RGBColor(r, g, b)
            
            elif element.name == 'p':
                # Handle paragraphs
                text = element.get_text().strip()
                if text:
                    p = doc.add_paragraph(text)
                    p.style = 'Normal'
            
            elif element.name in ['ul', 'ol']:
                # Handle lists
                list_id = id(element)
                list_type = element.name
                
                if list_type == 'ol':
                    ol_counters[list_id] = 1
                
                list_info = {
                    'type': list_type,
                    'id': list_id,
                    'level': len(list_stack)
                }
                list_stack.append(list_info)
                
                # Process list items
                for child in element.children:
                    if hasattr(child, 'name'):
                        process_element(child, list_info)
                
                list_stack.pop()
                
                # Clean up counter
                if list_type == 'ol' and list_id in ol_counters:
                    del ol_counters[list_id]
            
            elif element.name == 'li' and parent_list_info:
                # Handle list items
                text = element.get_text().strip()
                if text:
                    p = doc.add_paragraph(text)
                    
                    # Set appropriate list style
                    try:
                        if parent_list_info['type'] == 'ul':
                            p.style = 'List Bullet'
                        else:
                            p.style = 'List Number'
                    except KeyError:
                        # Fallback: format manually
                        p.style = 'Normal'
                        indent_level = parent_list_info['level']
                        p.paragraph_format.left_indent = Inches(0.5 * (indent_level + 1))
                        p.paragraph_format.first_line_indent = Inches(-0.25)
                        
                        if parent_list_info['type'] == 'ul':
                            # Add bullet
                            p.text = '• ' + p.text
                        else:
                            # Add number
                            list_id = parent_list_info['id']
                            num = ol_counters.get(list_id, 1)
                            p.text = f'{num}. ' + p.text
                            ol_counters[list_id] = num + 1
            
            elif element.name == 'table':
                # Handle tables
                rows = element.find_all('tr')
                if rows:
                    # Count columns
                    max_cols = max(len(row.find_all(['td', 'th'])) for row in rows)
                    
                    # Create table
                    table = doc.add_table(rows=0, cols=max_cols)
                    
                    # Add rows
                    for row in rows:
                        cells = row.find_all(['td', 'th'])
                        doc_row = table.add_row()
                        for i, cell in enumerate(cells):
                            if i < max_cols:
                                doc_row.cells[i].text = cell.get_text().strip()
            
            elif element.name == 'br':
                # Handle line breaks
                doc.add_paragraph()
            
            elif element.name == 'hr':
                # Handle horizontal rules
                p = doc.add_paragraph('_' * 50)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            elif element.name in ['em', 'i']:
                # Handle italic text (would need parent context for inline)
                pass
            
            elif element.name in ['strong', 'b']:
                # Handle bold text (would need parent context for inline)
                pass
            
            # Process child elements
            if element.name not in ['ul', 'ol', 'li']:
                for child in element.children:
                    if hasattr(child, 'name'):
                        process_element(child, parent_list_info)
        
        # Process all top-level elements
        for element in soup.body.children if soup.body else soup.children:
            if hasattr(element, 'name'):
                process_element(element)
    
    def replace_placeholders(self, template_path, placeholders, output_filepath=None, 
                           convert_markdown=True, preserve_template_styles=True):
        """
        Replace placeholders in a Word template with provided values.
        
        Args:
            template_path (str): Path to the template docx file
            placeholders (dict): Dictionary of placeholder:value pairs
                                e.g., {'company_name': 'DXE LTD', 'summary': '# Summary\nThis is markdown'}
            output_filepath (str, optional): Path to save the output file
            convert_markdown (bool): Whether to convert markdown values to formatted text
            preserve_template_styles (bool): Whether to preserve the template's existing styles
            
        Returns:
            BytesIO or bool: BytesIO object if output_filepath is None, else True/False
        """
        try:
            # Validate inputs
            if not os.path.exists(template_path):
                raise FileNotFoundError(f"Template file not found: {template_path}")
            
            if not isinstance(placeholders, dict):
                raise TypeError("placeholders must be a dictionary")
            
            # Load the template
            doc = Document(template_path)
            
            # Ensure required styles exist
            self._ensure_required_styles(doc)
            
            # Process each placeholder
            for placeholder, value in placeholders.items():
                # Format placeholder with braces if not already present
                if not placeholder.startswith('{'):
                    placeholder = '{' + placeholder + '}'
                
                # Handle markdown conversion
                if convert_markdown and self._is_likely_markdown(str(value)):
                    # Replace with formatted markdown
                    if not self._replace_placeholder_with_markdown(doc, placeholder, str(value)):
                        # If complex replacement failed, do simple replacement
                        self._simple_replace_all(doc, placeholder, str(value))
                else:
                    # Simple text replacement
                    self._simple_replace_all(doc, placeholder, str(value))
            
            # Apply custom table formatting if needed and not preserving styles
            if self.custom_table_style and not preserve_template_styles:
                self.format_table_with_custom_style(doc)
            
            # Save or return
            if output_filepath:
                # Ensure directory exists
                os.makedirs(os.path.dirname(os.path.abspath(output_filepath)), exist_ok=True)
                doc.save(output_filepath)
                return True
            else:
                docx_buffer = io.BytesIO()
                doc.save(docx_buffer)
                docx_buffer.seek(0)
                return docx_buffer
                
        except Exception as e:
            self.logger.exception(f"Error in replacing placeholders: {e}")
            if output_filepath:
                return False
            else:
                raise
    
    def _simple_replace_all(self, doc, find_text, replace_text):
        """
        Simple replacement of text throughout the document.
        
        Args:
            doc: Document object
            find_text: Text to find
            replace_text: Text to replace with
        """
        # Replace in paragraphs
        for paragraph in doc.paragraphs:
            if find_text in paragraph.text:
                self._find_and_replace_in_paragraph(paragraph, find_text, replace_text)
        
        # Replace in tables
        self._find_and_replace_in_tables(doc, find_text, replace_text)
        
        # Replace in headers and footers
        for section in doc.sections:
            # Headers
            for header in [section.header, section.first_page_header, section.even_page_header]:
                if header:
                    for paragraph in header.paragraphs:
                        if find_text in paragraph.text:
                            self._find_and_replace_in_paragraph(paragraph, find_text, replace_text)
            
            # Footers
            for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
                if footer:
                    for paragraph in footer.paragraphs:
                        if find_text in paragraph.text:
                            self._find_and_replace_in_paragraph(paragraph, find_text, replace_text)
    
    def markdown_to_docx(self, markdown_text, output_filepath=None, reference_docx_path=None,
                        preserve_reference_styles=False, custom_table_style_override=None):
        """
        Convert markdown text to a formatted Word document.
        
        Args:
            markdown_text (str): The markdown text to convert
            output_filepath (str, optional): Path to save the output file
            reference_docx_path (str, optional): Path to a reference docx file to use as a template
            preserve_reference_styles (bool): If True, maintains styles from reference document
            custom_table_style_override (bool, optional): Override table styling behavior
                                                         None = use self.custom_table_style
                                                         True = force custom table style
                                                         False = no custom table style
            
        Returns:
            BytesIO or bool: BytesIO object if output_filepath is None, else True/False
        """
        try:
            # Validate inputs
            if not isinstance(markdown_text, str):
                raise TypeError("markdown_text must be a string")
            
            if output_filepath:
                output_dir = os.path.dirname(os.path.abspath(output_filepath))
                if output_dir:
                    os.makedirs(output_dir, exist_ok=True)
            
            # Create document from reference or new
            if reference_docx_path and os.path.exists(reference_docx_path):
                doc = Document(reference_docx_path)
                # If preserving styles, don't clear the document
                if not preserve_reference_styles:
                    # Clear existing content if not preserving
                    for element in doc.element.body:
                        doc.element.body.remove(element)
            else:
                doc = Document()
            
            # Ensure required styles exist
            self._ensure_required_styles(doc)
            
            # Set page orientation and margins only if not preserving reference styles
            if not preserve_reference_styles:
                section = doc.sections[0]
                if self.page_orientation.lower() == "landscape":
                    section.orientation = WD_ORIENT.LANDSCAPE
                    section.page_width, section.page_height = section.page_height, section.page_width
                
                section.top_margin = Inches(self.margins['top'])
                section.bottom_margin = Inches(self.margins['bottom'])
                section.left_margin = Inches(self.margins['left'])
                section.right_margin = Inches(self.margins['right'])
            
            # Convert Markdown to HTML
            html = mistune.create_markdown(plugins=['table'], escape=False)(markdown_text)
            
            # Ensure valid HTML
            if html and not html.strip().startswith('<'):
                html = f'<div>{html}</div>'
            if html and not html.strip().startswith('<!DOCTYPE') and not html.strip().startswith('<html'):
                html = f'<html><body>{html}</body></html>'
            
            # Try to use HtmlToDocx, but fall back to custom parser if it fails
            try:
                HtmlToDocx().add_html_to_document(html, doc)
            except Exception as e:
                self.logger.warning(f"HtmlToDocx failed: {e}. Using custom parser.")
                # Clear the document and use custom parser
                for element in doc.element.body:
                    doc.element.body.remove(element)
                self._custom_html_parser(html, doc)
            
            # Apply formatting only if not preserving reference styles
            if not preserve_reference_styles:
                # Apply custom table formatting if needed
                if self.custom_table_style:
                    self.format_table_with_custom_style(doc)
                
                # Set default font for paragraphs and adjust headings
                for p in doc.paragraphs:
                    for run in p.runs:
                        run.font.name = self.font_name
                        run.font.size = Pt(self.font_size)
                    
                    if p.style.name.startswith('Heading'):
                        try:
                            level = int(p.style.name[-1])
                        except ValueError:
                            level = 1
                        
                        if level <= 4:  # Support up to h4
                            for run in p.runs:
                                run.font.name = self.heading_font
                                run.font.size = Pt(self.heading_sizes.get(f'h{level}', self.font_size))
                                run.font.bold = True
                                
                                if self.heading_color:
                                    hex_color = self.heading_color.lstrip('#')
                                    r, g, b = int(hex_color[:2], 16), int(hex_color[2:4], 16), int(hex_color[4:6], 16)
                                    run.font.color.rgb = RGBColor(r, g, b)
            else:
                # When preserving reference styles, check if we should still apply custom table styling
                apply_table_style = custom_table_style_override if custom_table_style_override is not None else self.custom_table_style
                if apply_table_style:
                    self.format_table_with_custom_style(doc)
            
            # Return result
            if output_filepath:
                doc.save(output_filepath)
                return True
            else:
                docx_buffer = io.BytesIO()
                doc.save(docx_buffer)
                docx_buffer.seek(0)
                return docx_buffer
                
        except Exception as e:
            self.logger.exception(f"Error in converting Markdown to DOCX: {e}")
            if output_filepath:
                return False
            else:
                raise
    
    def markdown_to_memory(self, markdown_text, reference_docx_path=None, preserve_reference_styles=False):
        """
        Convert markdown text to a Word document and return as BytesIO object.
        
        Args:
            markdown_text (str): The markdown text to convert
            reference_docx_path (str, optional): Path to a reference docx file
            preserve_reference_styles (bool): Whether to preserve reference document styles
            
        Returns:
            BytesIO: BytesIO object containing the document
        """
        return self.markdown_to_docx(markdown_text, output_filepath=None, 
                                   reference_docx_path=reference_docx_path,
                                   preserve_reference_styles=preserve_reference_styles)
    
    # Setters for configuration
    def set_font(self, font_name, font_size=None):
        """Update the default font settings"""
        self.font_name = font_name
        if font_size is not None:
            self.font_size = font_size
    
    def set_heading_style(self, font=None, sizes=None, color=None):
        """Update heading style settings"""
        if font is not None:
            self.heading_font = font
        if sizes is not None:
            self.heading_sizes = sizes
        if color is not None:
            self.heading_color = color
    
    def set_page_orientation(self, orientation):
        """Set page orientation ('portrait' or 'landscape')"""
        if orientation.lower() in ['portrait', 'landscape']:
            self.page_orientation = orientation.lower()
        else:
            raise ValueError("Orientation must be 'portrait' or 'landscape'")
    
    def set_margins(self, margins):
        """Set page margins in inches"""
        if isinstance(margins, dict) and all(k in margins for k in ['top', 'bottom', 'left', 'right']):
            self.margins = margins
        else:
            raise ValueError("Margins must be a dict with keys: top, bottom, left, right")